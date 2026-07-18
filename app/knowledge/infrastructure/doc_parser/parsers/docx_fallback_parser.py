"""
RAG 文档解析器 — python-docx 兜底解析器。

当 Docling 不可用或解析失败时，使用 python-docx 解析 DOCX。
实现 iter_block_items 以保留段落和表格的原始顺序。
"""

from __future__ import annotations

import logging
from typing import List, Optional

from app.knowledge.infrastructure.doc_parser.config import ParserConfig
from app.knowledge.infrastructure.doc_parser.exceptions import DocumentParseError
from app.knowledge.infrastructure.doc_parser.models import PageMarkdown, ParsedMarkdownDocument
from app.knowledge.infrastructure.doc_parser.parsers.base import BaseDocumentParser

logger = logging.getLogger(__name__)


class DocxFallbackParser(BaseDocumentParser):
    """基于 python-docx 的 DOCX 兜底解析器。

    当 Docling 不可用时使用。支持：
    - 段落（标题、正文、列表）
    - 表格（转为 Markdown 表格）
    - 保留段落和表格的原始顺序
    """

    def __init__(self, config: Optional[ParserConfig] = None) -> None:
        """初始化 python-docx 兜底解析器。"""
        super().__init__(config)
        self.parser_name = "DocxFallbackParser"

    def parse(self, file_path: str, doc_id: str) -> ParsedMarkdownDocument:
        """解析 DOCX 文档为统一 Markdown 格式。

        Args:
            file_path: DOCX 文件路径。
            doc_id: 文档唯一标识。

        Returns:
            ParsedMarkdownDocument 实例。

        Raises:
            DocumentParseError: 解析失败时抛出。
        """
        self._validate_file(file_path, [".docx"])
        logger.info("[%s] 开始解析 DOCX: %s (doc_id=%s)", self.parser_name, file_path, doc_id)

        try:
            from docx import Document

            doc = Document(file_path)

            # 按顺序遍历段落和表格，生成 Markdown
            markdown_lines: List[str] = []
            table_count = 0

            for block in self._iter_block_items(doc):
                block_type, content = block

                if block_type == "paragraph":
                    md = self._paragraph_to_markdown(content)
                    if md:
                        markdown_lines.append(md)

                elif block_type == "table":
                    md = self._table_to_markdown(content)
                    if md:
                        markdown_lines.append(md)
                        table_count += 1

            markdown_text = "\n\n".join(markdown_lines)

            # 构建 metadata
            metadata = self._build_metadata(
                picture_count=0,
                table_count=table_count,
                page_count=0,
                warnings=["使用 python-docx 兜底解析器，图片描述功能不可用。"],
            )

            # 构建 page_markdown_list
            page_markdown_list = [
                PageMarkdown(
                    page_number=1,
                    markdown=markdown_text,
                    metadata={"source": "python-docx"},
                )
            ]

            logger.info("[%s] DOCX 解析完成: tables=%d", self.parser_name, table_count)

            return ParsedMarkdownDocument(
                doc_id=doc_id,
                source_file=file_path,
                markdown=markdown_text,
                page_markdown_list=page_markdown_list,
                metadata=metadata,
            )

        except Exception as e:
            raise DocumentParseError(
                str(e), file_path=file_path, parser_name=self.parser_name
            ) from e

    @staticmethod
    def _iter_block_items(doc):
        """按文档顺序遍历段落和表格。

        python-docx 的 doc.paragraphs 和 doc.tables 是分开的，
        无法保留原始顺序。此方法通过遍历 document.element.body
        来保持段落和表格的正确顺序。

        Args:
            doc: python-docx Document 对象。

        Yields:
            ("paragraph", Paragraph) 或 ("table", Table) 元组。
        """
        from docx.oxml.ns import qn

        body = doc.element.body
        for child in body:
            tag = child.tag
            # 段落标签
            if tag == qn("w:p"):
                yield ("paragraph", _wrap_paragraph(doc, child))
            # 表格标签
            elif tag == qn("w:tbl"):
                yield ("table", _wrap_table(doc, child))

    def _paragraph_to_markdown(self, paragraph) -> str:
        """将段落转为 Markdown。

        处理标题（Heading 1-6）、列表项、普通段落。

        Args:
            paragraph: python-docx Paragraph 对象。

        Returns:
            Markdown 格式文本。
        """
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = (paragraph.style.name or "").lower()

        # 标题
        if "heading" in style_name:
            # 提取标题级别
            try:
                level = int(style_name.replace("heading", "").strip())
                level = min(max(level, 1), 6)
            except ValueError:
                level = 1
            return f"{'#' * level} {text}"

        # 列表项
        if "list" in style_name:
            # 检查是否有序列表
            if "number" in style_name:
                return f"1. {text}"
            return f"- {text}"

        # 普通段落
        return text

    def _table_to_markdown(self, table) -> str:
        """将表格转为 Markdown 表格格式。

        Args:
            table: python-docx Table 对象。

        Returns:
            Markdown 表格文本。
        """
        rows_data: List[List[str]] = []

        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return ""

        # 确定列数（取最大列数）
        max_cols = max(len(row) for row in rows_data)
        # 补齐列数
        for row in rows_data:
            while len(row) < max_cols:
                row.append("")

        lines: List[str] = []

        # 第一行作为表头
        header = rows_data[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        # 数据行
        for row in rows_data[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)


def _wrap_paragraph(doc, element):
    """将 XML element 包装为 Paragraph 对象。

    Args:
        doc: Document 对象。
        element: lxml Element。

    Returns:
        Paragraph 对象。
    """
    from docx.text.paragraph import Paragraph
    return Paragraph(element, doc)


def _wrap_table(doc, element):
    """将 XML element 包装为 Table 对象。

    Args:
        doc: Document 对象。
        element: lxml Element。

    Returns:
        Table 对象。
    """
    from docx.table import Table
    return Table(element, doc)
